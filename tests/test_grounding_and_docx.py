import io
import unittest

from docx import Document

from prompts.templates import (
    RESUME_WRITER_SYSTEM_PROMPT,
    build_achievement_prompt,
    build_achievement_refinement_prompt,
    build_cover_letter_refinement_prompt,
    build_grounded_resume_prompt,
    build_resume_refinement_prompt,
)
from utils.docx_builder import (
    make_docx_from_text,
    validate_docx_roundtrip,
)


class GroundingAndDocumentTests(unittest.TestCase):
    def test_grounding_prompt_forbids_invented_metrics(self):
        prompt = build_achievement_prompt(
            count=3,
            key_requirements="- SQL",
            job_title="Data Analyst",
            achievement_evidence=(
                "[E001] section=experience | verification=source_explicit | "
                "Prepared weekly SQL reports."
            ),
        )
        combined = (RESUME_WRITER_SYSTEM_PROMPT + prompt).lower()
        self.assertIn("never add a metric", combined)
        self.assertIn("[metric needed:", combined)
        self.assertIn("verified candidate evidence", combined)
        self.assertIn("evidence: e###", combined)
        self.assertIn("prepared weekly sql reports", combined)

    def test_achievement_prompt_refuses_untagged_or_jd_only_input(self):
        with self.assertRaisesRegex(ValueError, "requires experience"):
            build_achievement_prompt(
                count=3,
                key_requirements="- SQL and AWS",
                job_title="Data Analyst",
                achievement_evidence="Prepared weekly reports.",
            )

    def test_full_resume_prompt_requires_traceable_candidate_evidence(self):
        with self.assertRaisesRegex(ValueError, "requires candidate evidence"):
            build_grounded_resume_prompt(
                fields="Data Analytics",
                job_description="SQL required. Improve reporting by 30%.",
                candidate_evidence="Candidate has reporting experience.",
            )

    def test_full_resume_prompt_makes_jd_context_only(self):
        prompt = build_grounded_resume_prompt(
            fields="Data Analytics",
            job_description="SQL required. Improve reporting by 30%.",
            candidate_evidence=(
                "[E001] section=experience | verification=source_explicit | "
                "Prepared weekly operational reports."
            ),
        )
        lowered = prompt.lower()
        self.assertIn("candidate evidence is the only source", lowered)
        self.assertIn("never estimate, round, extrapolate", lowered)
        self.assertIn("job description (requirements only; never candidate evidence)", lowered)
        self.assertIn("evidence: e###", lowered)
        self.assertIn("jd match: r###", lowered)
        self.assertNotIn("reasonable to infer", lowered)
        self.assertNotIn("quantify everything", lowered)

    def test_second_pass_prompts_require_evidence_and_target_genericity(self):
        evidence = (
            "[E001] section=experience | Prepared weekly SQL reports.\n"
            "[R001] required | direct | evidence=E001 | Prepare reports"
        )
        resume_review = build_resume_refinement_prompt(
            draft=(
                "- Prepared weekly SQL reports.\n"
                '  Evidence: E001 — "Prepared weekly SQL reports."\n'
                '  JD Match: R001 — "Prepare reports"'
            ),
            candidate_evidence=evidence,
        ).lower()
        achievement_review = build_achievement_refinement_prompt(
            draft="- Prepared weekly SQL reports.",
            key_requirements="- Prepare reports",
            achievement_evidence=evidence,
            count=3,
        ).lower()
        cover_review = build_cover_letter_refinement_prompt(
            draft="I am writing to apply.",
            tone="Confident",
            candidate_evidence=evidence,
            requirement_context="[R001] direct | Prepare reports",
        ).lower()
        self.assertIn("requirement fit", resume_review)
        self.assertIn("generic filler", resume_review)
        self.assertIn("prefer evidence that genuinely addresses", achievement_review)
        self.assertIn("remove generic enthusiasm", cover_review)

    def test_docx_uses_candidate_name_not_target_title(self):
        payload = make_docx_from_text(
            "PROFESSIONAL SUMMARY\nAnalyst with verified reporting experience.",
            name="Jane Doe",
        )
        document = Document(io.BytesIO(payload))
        text = "\n".join(paragraph.text for paragraph in document.paragraphs)
        self.assertIn("Jane Doe", text)
        self.assertNotIn("Data Analyst", text)

    def test_docx_roundtrip_preserves_ats_readable_resume_structure(self):
        source = """\
Jane Doe
jane@example.com | +254 700 000 000

PROFESSIONAL SUMMARY
Data analyst with verified reporting experience.

CORE SKILLS
SQL, Python, Power BI

PROFESSIONAL EXPERIENCE
Data Analyst | Acme Ltd | Nairobi | 2022 - Present
- Built weekly SQL reports for management.

EDUCATION
BSc Statistics | Example University | 2021
"""
        payload = make_docx_from_text(source)
        report = validate_docx_roundtrip(source, payload)
        self.assertTrue(report.is_safe)
        self.assertGreaterEqual(report.retention_rate, 0.95)
        self.assertFalse(report.missing_sections)
        self.assertTrue(report.contact_retained)
        self.assertEqual(report.roles_retained, report.roles_expected)

    def test_docx_roundtrip_detects_lost_sections_and_content(self):
        source = """\
Jane Doe
jane@example.com
EXPERIENCE
Data Analyst | Acme Ltd | 2022 - Present
- Built weekly SQL reports.
EDUCATION
BSc Statistics | Example University | 2021
"""
        broken = Document()
        broken.add_paragraph("Jane Doe")
        payload = io.BytesIO()
        broken.save(payload)
        report = validate_docx_roundtrip(source, payload.getvalue())
        self.assertFalse(report.is_safe)
        self.assertLess(report.retention_rate, 0.95)
        self.assertIn("experience", report.missing_sections)
        self.assertTrue(report.issues)

    def test_docx_roundtrip_flags_table_based_layout(self):
        source = "Jane Doe\njane@example.com"
        document = Document()
        table = document.add_table(rows=1, cols=1)
        table.cell(0, 0).text = source
        payload = io.BytesIO()
        document.save(payload)
        report = validate_docx_roundtrip(source, payload.getvalue())
        self.assertEqual(report.table_count, 1)
        self.assertFalse(report.is_safe)
        self.assertTrue(any("table" in issue.lower() for issue in report.issues))


if __name__ == "__main__":
    unittest.main()
