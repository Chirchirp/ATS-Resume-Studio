"""
utils/docx_builder.py — Convert resume / cover-letter text to a formatted DOCX file.
"""

import io
import re
from collections import Counter
from dataclasses import dataclass

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from utils.ats_engine import canonical_resume_section, extract_resume_profile
from utils.text_processing import BULLET_PREFIX_RE, normalize_resume_structure

BOLD_PATTERN = re.compile(r"\*\*(.+?)\*\*")
NAVY = "17365D"
TEAL = "1B7F79"
BODY = "263645"
MUTED = "536779"


@dataclass(frozen=True)
class DocxParseabilityReport:
    score: int
    retention_rate: float
    source_word_count: int
    extracted_word_count: int
    expected_sections: tuple[str, ...]
    extracted_sections: tuple[str, ...]
    missing_sections: tuple[str, ...]
    contact_retained: bool
    roles_retained: int
    roles_expected: int
    bullets_retained: int
    bullets_expected: int
    table_count: int
    multi_column_sections: int
    drawing_count: int
    header_footer_text_present: bool
    issues: tuple[str, ...]
    extracted_text: str

    @property
    def is_safe(self) -> bool:
        return (
            self.score >= 90
            and not self.missing_sections
            and self.table_count == 0
            and self.multi_column_sections == 0
            and self.drawing_count == 0
            and not self.header_footer_text_present
            and self.bullets_retained >= self.bullets_expected
        )


def _add_runs_with_bold(paragraph, text: str):
    """Add text runs to a paragraph, honouring **bold** markers."""
    if not text:
        return
    last = 0
    found = False
    for m in BOLD_PATTERN.finditer(text):
        found = True
        pre = text[last : m.start()]
        if pre:
            paragraph.add_run(pre)
        paragraph.add_run(m.group(1)).bold = True
        last = m.end()
    remaining = text[last:]
    if not found:
        remaining = re.sub(r"\*\*", "", remaining)
    if remaining:
        paragraph.add_run(remaining)


def _add_horizontal_line(doc):
    """Insert a thin horizontal rule paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "CCCCCC")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _set_style_font(style, name: str, size: float, color: str = BODY):
    style.font.name = name
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)
    style._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), name)
    style._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), name)


def _set_bottom_border(paragraph, color: str = TEAL, size: str = "8"):
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), color)
    borders.append(bottom)


def _section_key(line: str) -> str | None:
    return canonical_resume_section(line)


def _next_numbering_id(numbering, tag: str, attribute: str) -> int:
    values = []
    for element in numbering.findall(qn(tag)):
        value = element.get(qn(attribute))
        if value and value.isdigit():
            values.append(int(value))
    return max(values, default=0) + 1


def _create_resume_bullet_numbering(doc: Document) -> int:
    """Create a real Word bullet definition with explicit ATS-resume geometry."""
    numbering = doc.part.numbering_part.element
    abstract_id = _next_numbering_id(
        numbering, "w:abstractNum", "w:abstractNumId"
    )
    num_id = _next_numbering_id(numbering, "w:num", "w:numId")

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    number_format = OxmlElement("w:numFmt")
    number_format.set(qn("w:val"), "bullet")
    level.append(number_format)
    level_text = OxmlElement("w:lvlText")
    level_text.set(qn("w:val"), "•")
    level.append(level_text)
    justification = OxmlElement("w:lvlJc")
    justification.set(qn("w:val"), "left")
    level.append(justification)

    paragraph_properties = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    paragraph_properties.append(tabs)
    indent = OxmlElement("w:ind")
    indent.set(qn("w:left"), "540")
    indent.set(qn("w:hanging"), "280")
    paragraph_properties.append(indent)
    level.append(paragraph_properties)

    run_properties = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), "Arial")
    fonts.set(qn("w:hAnsi"), "Arial")
    run_properties.append(fonts)
    level.append(run_properties)
    abstract.append(level)
    numbering.append(abstract)

    number = OxmlElement("w:num")
    number.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    number.append(abstract_ref)
    numbering.append(number)
    return num_id


def _apply_resume_bullet_numbering(paragraph, num_id: int) -> None:
    properties = paragraph._p.get_or_add_pPr()
    existing = properties.find(qn("w:numPr"))
    if existing is not None:
        properties.remove(existing)
    numbering = OxmlElement("w:numPr")
    level = OxmlElement("w:ilvl")
    level.set(qn("w:val"), "0")
    number = OxmlElement("w:numId")
    number.set(qn("w:val"), str(num_id))
    numbering.extend([level, number])
    properties.append(numbering)


def _style_pipe_record(paragraph, line: str):
    """Style a pipe-delimited record without changing its ATS extraction order."""
    parts = [re.sub(r"\*\*", "", part.strip()) for part in line.split("|")]
    for index, part in enumerate(parts):
        if index:
            separator = paragraph.add_run("  |  ")
            separator.font.color.rgb = RGBColor.from_string(TEAL)
        run = paragraph.add_run(part)
        run.bold = index == 0
        run.font.color.rgb = RGBColor.from_string(NAVY if index == 0 else BODY)


def make_docx_from_text(text: str, name: str = "") -> bytes:
    """
    Convert plain / lightly-marked-up resume text into a polished DOCX.

    Returns raw bytes suitable for st.download_button.
    """
    text = normalize_resume_structure(text)
    doc = Document()
    doc.core_properties.title = "ATS-Optimized Resume"
    doc.core_properties.subject = "Single-column ATS-readable professional resume"

    # Named design override: compact_reference_guide -> ats_resume_single_column.
    # The narrower 0.72" margins keep two-page resumes concise while all content
    # remains in the normal document body and reading order.
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.68)
    section.bottom_margin = Inches(0.68)
    section.left_margin = Inches(0.72)
    section.right_margin = Inches(0.72)

    style = doc.styles["Normal"]
    _set_style_font(style, "Arial", 10.25)
    pf = style.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf.line_spacing = 1.1
    pf.space_after = Pt(3)
    pf.space_before = Pt(0)

    heading_style = doc.styles["Heading 1"]
    _set_style_font(heading_style, "Arial", 11.5, NAVY)
    heading_style.font.bold = True
    heading_style.font.all_caps = True
    heading_style.paragraph_format.space_before = Pt(10)
    heading_style.paragraph_format.space_after = Pt(4)
    heading_style.paragraph_format.keep_with_next = True

    bullet_style = doc.styles["List Bullet"]
    _set_style_font(bullet_style, "Arial", 10.25)
    bullet_style.paragraph_format.left_indent = Inches(0.375)
    bullet_style.paragraph_format.first_line_indent = Inches(-0.194)
    bullet_style.paragraph_format.space_before = Pt(0)
    bullet_style.paragraph_format.space_after = Pt(3)
    bullet_style.paragraph_format.line_spacing = 1.1
    bullet_num_id = _create_resume_bullet_numbering(doc)

    parsed_profile = extract_resume_profile(text)
    candidate_name = (name or parsed_profile.candidate_name).strip()
    name_written = False

    lines = text.splitlines()
    i = 0
    in_core_skills = False

    while i < len(lines):
        line = lines[i].strip()

        if not line:
            i += 1
            continue

        normalized_plain = re.sub(r"\*\*", "", line).strip()
        if (
            candidate_name
            and not name_written
            and normalized_plain.casefold() == candidate_name.casefold()
        ):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(candidate_name)
            run.bold = True
            run.font.name = "Arial"
            run.font.size = Pt(21)
            run.font.color.rgb = RGBColor.from_string(NAVY)
            name_written = True
            i += 1
            continue

        if candidate_name and not name_written:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(candidate_name)
            run.bold = True
            run.font.name = "Arial"
            run.font.size = Pt(21)
            run.font.color.rgb = RGBColor.from_string(NAVY)
            name_written = True

        # Section header
        section_key = _section_key(line)
        if section_key:
            in_core_skills = section_key == "skills"
            heading_text = re.sub(r"[#\*_\-]{2,}", "", line.strip().rstrip(":")).strip()
            h = doc.add_heading(heading_text, level=1)
            h.alignment = WD_ALIGN_PARAGRAPH.LEFT
            _set_bottom_border(h)
            i += 1
            continue

        # Skill category line (no bullets)
        if (
            ("–" in line or "—" in line or (":" in line and len(line.split(":")[0].split()) <= 6))
            and not line.startswith("[")
            and "|" not in line
        ):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.space_before = Pt(0)

            if "–" in line:
                cat, *rest = line.split("–", 1)
                sep = " – "
            elif "—" in line:
                cat, *rest = line.split("—", 1)
                sep = " — "
            else:
                cat, *rest = line.split(":", 1)
                sep = ": "

            p.add_run(re.sub(r"\*\*", "", cat.strip())).bold = True
            if rest:
                p.add_run(sep + re.sub(r"\*\*", "", rest[0].strip()))
            i += 1
            continue

        # Bullet point
        if not in_core_skills and BULLET_PREFIX_RE.match(line):
            bullet_text = BULLET_PREFIX_RE.sub("", line, count=1).strip()
            p = doc.add_paragraph(style="List Bullet")
            _apply_resume_bullet_numbering(p, bullet_num_id)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.widow_control = True
            _add_runs_with_bold(p, bullet_text)
            i += 1
            continue

        # Contact info remains in the body (not a header) for ATS reliability.
        if any(kw in line.lower() for kw in ["@", "phone", "email", "linkedin", "github", "number"]):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_after = Pt(5)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.keep_with_next = True
            _add_runs_with_bold(p, line)
            for run in p.runs:
                run.font.size = Pt(9.25)
                run.font.color.rgb = RGBColor.from_string(MUTED)
            i += 1
            continue

        # Pipe-delimited role or qualification record.
        if "|" in line and not line.startswith("["):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.keep_with_next = True
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            _style_pipe_record(p, line)
            i += 1
            continue

        # Generic paragraph
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.space_before = Pt(0)
        _add_runs_with_bold(p, line)
        i += 1

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


def extract_text_from_docx(payload: bytes) -> str:
    """Extract paragraphs and table cells as a basic ATS-style round trip."""
    document = Document(io.BytesIO(payload))
    lines: list[str] = []
    for paragraph in document.paragraphs:
        text = re.sub(r"\s+", " ", paragraph.text).strip()
        if not text:
            continue
        style_name = str(getattr(paragraph.style, "name", "") or "").lower()
        has_numbering = bool(
            paragraph._p.pPr is not None
            and paragraph._p.pPr.find(qn("w:numPr")) is not None
        )
        if (
            ("list" in style_name or has_numbering)
            and not BULLET_PREFIX_RE.match(text)
        ):
            text = "- " + text
        lines.append(text)
    for table in document.tables:
        for row in table.rows:
            values = [
                re.sub(r"\s+", " ", cell.text).strip()
                for cell in row.cells
                if cell.text.strip()
            ]
            if values:
                lines.append(" | ".join(values))
    return "\n".join(lines)


def _comparison_tokens(text: str) -> Counter:
    cleaned = re.sub(r"[*_#•]", " ", text.lower())
    return Counter(re.findall(r"[a-z0-9][a-z0-9+#./%'-]*", cleaned))


def _normalized_line(text: str) -> str:
    return re.sub(r"[^a-z0-9]+", " ", text.lower()).strip()


def validate_docx_roundtrip(
    source_text: str,
    payload: bytes,
) -> DocxParseabilityReport:
    """Verify that a generated DOCX retains ATS-readable candidate content."""
    source_text = normalize_resume_structure(source_text)
    extracted = extract_text_from_docx(payload)
    document = Document(io.BytesIO(payload))
    source_tokens = _comparison_tokens(source_text)
    extracted_tokens = _comparison_tokens(extracted)
    source_count = sum(source_tokens.values())
    overlap = sum(
        min(count, extracted_tokens.get(token, 0))
        for token, count in source_tokens.items()
    )
    retention = overlap / source_count if source_count else 0.0

    source_profile = extract_resume_profile(source_text)
    extracted_profile = extract_resume_profile(extracted)
    expected_sections = tuple(source_profile.section_order)
    extracted_sections = tuple(extracted_profile.section_order)
    missing_sections = tuple(
        section for section in expected_sections if section not in extracted_sections
    )

    expected_emails = set(source_profile.contact.emails)
    extracted_emails = set(extracted_profile.contact.emails)
    expected_phones = {
        re.sub(r"\D", "", value) for value in source_profile.contact.phones
    }
    extracted_phones = {
        re.sub(r"\D", "", value) for value in extracted_profile.contact.phones
    }
    contact_expected = bool(expected_emails or expected_phones)
    contact_retained = (
        (not expected_emails or expected_emails <= extracted_emails)
        and (not expected_phones or expected_phones <= extracted_phones)
    )

    extracted_normalized = _normalized_line(extracted)
    roles_expected = len(source_profile.roles)
    roles_retained = sum(
        _normalized_line(role.header) in extracted_normalized
        for role in source_profile.roles
    )
    bullets_expected = sum(
        bool(re.match(r"^\s*[-*•▪◦●–—]\s+", line))
        for line in source_text.splitlines()
    )
    bullets_retained = sum(
        bool(re.match(r"^\s*[-*•▪◦●–—]\s+", line))
        for line in extracted.splitlines()
    )

    section_rate = (
        (len(expected_sections) - len(missing_sections)) / len(expected_sections)
        if expected_sections
        else 1.0
    )
    role_rate = (
        roles_retained / roles_expected if roles_expected else 1.0
    )
    contact_rate = 1.0 if contact_retained or not contact_expected else 0.0
    table_count = len(document.tables)
    multi_column_sections = 0
    for section in document.sections:
        columns = section._sectPr.xpath("./w:cols")
        if columns:
            value = columns[0].get(qn("w:num"))
            if value and value.isdigit() and int(value) > 1:
                multi_column_sections += 1
    drawing_count = len(
        document.element.xpath(".//w:drawing | .//w:pict")
    )
    header_footer_text_present = any(
        paragraph.text.strip()
        for section in document.sections
        for container in (section.header, section.footer)
        for paragraph in container.paragraphs
    )

    score = round(
        min(1.0, retention) * 60
        + section_rate * 20
        + contact_rate * 10
        + role_rate * 10
    )
    score -= min(20, table_count * 10)
    score -= min(30, multi_column_sections * 15)
    score -= min(15, drawing_count * 5)
    if header_footer_text_present:
        score -= 5

    issues: list[str] = []
    if retention < 0.95:
        issues.append(
            f"Only {retention:.0%} of source tokens survived DOCX extraction."
        )
    if missing_sections:
        issues.append(
            "Missing extracted sections: " + ", ".join(missing_sections)
        )
    if contact_expected and not contact_retained:
        issues.append("One or more source contact fields did not survive extraction.")
    if roles_retained < roles_expected:
        issues.append(
            f"Only {roles_retained} of {roles_expected} role headers survived extraction."
        )
    if bullets_retained < bullets_expected:
        issues.append(
            f"Only {bullets_retained} of {bullets_expected} bullet boundaries "
            "survived DOCX extraction."
        )
    if table_count:
        issues.append(
            f"Detected {table_count} table(s); table reading order can vary across parsers."
        )
    if multi_column_sections:
        issues.append(
            f"Detected {multi_column_sections} multi-column section(s), which can "
            "interleave extracted text."
        )
    if drawing_count:
        issues.append(
            f"Detected {drawing_count} drawing/image object(s); text inside graphics "
            "is not treated as readable evidence."
        )
    if header_footer_text_present:
        issues.append(
            "Detected header/footer text; critical contact or resume content should "
            "remain in the document body."
        )

    return DocxParseabilityReport(
        score=max(0, min(100, score)),
        retention_rate=retention,
        source_word_count=source_count,
        extracted_word_count=sum(extracted_tokens.values()),
        expected_sections=expected_sections,
        extracted_sections=extracted_sections,
        missing_sections=missing_sections,
        contact_retained=contact_retained,
        roles_retained=roles_retained,
        roles_expected=roles_expected,
        bullets_retained=bullets_retained,
        bullets_expected=bullets_expected,
        table_count=table_count,
        multi_column_sections=multi_column_sections,
        drawing_count=drawing_count,
        header_footer_text_present=header_footer_text_present,
        issues=tuple(issues),
        extracted_text=extracted,
    )
